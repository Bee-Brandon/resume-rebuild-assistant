"""Generate the default.docx template with docxtpl Jinja2 tokens.

Run once:  python create_template.py
Produces:  templates/default.docx
"""
from pathlib import Path

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH


def _add_section_heading(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.keep_with_next = True
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(14)
    run.font.name = "Calibri"
    run.font.color.rgb = RGBColor(0x1A, 0x1A, 0x2E)
    # Bottom border on the paragraph
    from docx.oxml.ns import qn
    pPr = p._element.get_or_add_pPr()
    pBdr = pPr.makeelement(qn("w:pBdr"), {})
    bottom = pBdr.makeelement(qn("w:bottom"), {
        qn("w:val"): "single",
        qn("w:sz"): "4",
        qn("w:space"): "1",
        qn("w:color"): "CCCCCC",
    })
    pBdr.append(bottom)
    pPr.append(pBdr)
    return p


def _set_default_font(doc):
    style = doc.styles["Normal"]
    font = style.font
    font.name = "Calibri"
    font.size = Pt(11)
    font.color.rgb = RGBColor(0x1A, 0x1A, 0x1A)
    pf = style.paragraph_format
    pf.space_before = Pt(0)
    pf.space_after = Pt(2)
    pf.line_spacing = 1.15


def create_template(output_path: Path):
    doc = Document()
    _set_default_font(doc)

    # Set 0.5 inch margins
    for section in doc.sections:
        section.top_margin = Inches(0.5)
        section.bottom_margin = Inches(0.5)
        section.left_margin = Inches(0.5)
        section.right_margin = Inches(0.5)

    # ── HEADER: Name ──
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run("{{contact.name}}")
    run.bold = True
    run.font.size = Pt(20)
    run.font.name = "Calibri"

    # Contact line
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run("{{contact_line}}")
    run.font.size = Pt(10)
    run.font.name = "Calibri"
    run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

    # LinkedIn (conditional)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run("{% if contact.linkedin %}{{contact.linkedin}}{% endif %}")
    run.font.size = Pt(10)
    run.font.name = "Calibri"
    run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

    # ── SUMMARY ──
    p = doc.add_paragraph()
    run = p.add_run("{%p if has_summary %}")
    run.font.size = Pt(1)
    run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

    _add_section_heading(doc, "Summary")

    p = doc.add_paragraph()
    run = p.add_run("{{summary}}")
    run.font.name = "Calibri"
    run.font.size = Pt(11)

    p = doc.add_paragraph()
    run = p.add_run("{%p endif %}")
    run.font.size = Pt(1)
    run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

    # ── SKILLS ──
    p = doc.add_paragraph()
    run = p.add_run("{%p if has_skills %}")
    run.font.size = Pt(1)
    run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

    _add_section_heading(doc, "Skills")

    p = doc.add_paragraph()
    run = p.add_run("{{skills_text}}")
    run.font.name = "Calibri"
    run.font.size = Pt(11)

    p = doc.add_paragraph()
    run = p.add_run("{%p endif %}")
    run.font.size = Pt(1)
    run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

    # ── EXPERIENCE ──
    p = doc.add_paragraph()
    run = p.add_run("{%p if has_work %}")
    run.font.size = Pt(1)
    run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

    _add_section_heading(doc, "Experience")

    # Loop start
    p = doc.add_paragraph()
    run = p.add_run("{%p for job in work_history %}")
    run.font.size = Pt(1)
    run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

    # Job title + employer
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(1)
    run = p.add_run("{{job.role}}")
    run.bold = True
    run.font.name = "Calibri"
    run.font.size = Pt(11)
    run2 = p.add_run(" \u2014 {{job.employer}}")
    run2.font.name = "Calibri"
    run2.font.size = Pt(11)

    # Dates + location
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run("{{job.start_date}} \u2013 {{job.end_date}}")
    run.italic = True
    run.font.name = "Calibri"
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
    run2 = p.add_run("{% if job.location %} | {{job.location}}{% endif %}")
    run2.italic = True
    run2.font.name = "Calibri"
    run2.font.size = Pt(10)
    run2.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

    # Bullets loop
    p = doc.add_paragraph()
    run = p.add_run("{%p for bullet in job.bullets %}")
    run.font.size = Pt(1)
    run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(1)
    p.paragraph_format.left_indent = Inches(0.25)
    run = p.add_run("{{bullet}}")
    run.font.name = "Calibri"
    run.font.size = Pt(11)

    p = doc.add_paragraph()
    run = p.add_run("{%p endfor %}")
    run.font.size = Pt(1)
    run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

    # End job loop
    p = doc.add_paragraph()
    run = p.add_run("{%p endfor %}")
    run.font.size = Pt(1)
    run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

    p = doc.add_paragraph()
    run = p.add_run("{%p endif %}")
    run.font.size = Pt(1)
    run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

    # ── EDUCATION ──
    p = doc.add_paragraph()
    run = p.add_run("{%p if has_education %}")
    run.font.size = Pt(1)
    run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

    _add_section_heading(doc, "Education")

    p = doc.add_paragraph()
    run = p.add_run("{%p for edu in education %}")
    run.font.size = Pt(1)
    run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run("{{edu.institution}}")
    run.bold = True
    run.font.name = "Calibri"
    run.font.size = Pt(11)
    run2 = p.add_run(
        "{% if edu.degree %} \u2014 {{edu.degree}}"
        "{% if edu.field %} in {{edu.field}}{% endif %}"
        "{% endif %}"
        "{% if edu.graduation_year %}, {{edu.graduation_year}}{% endif %}"
    )
    run2.font.name = "Calibri"
    run2.font.size = Pt(11)

    # Honors
    p = doc.add_paragraph()
    run = p.add_run("{% if edu.honors %}{{edu.honors}}{% endif %}")
    run.italic = True
    run.font.name = "Calibri"
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

    p = doc.add_paragraph()
    run = p.add_run("{%p endfor %}")
    run.font.size = Pt(1)
    run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

    p = doc.add_paragraph()
    run = p.add_run("{%p endif %}")
    run.font.size = Pt(1)
    run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

    # ── CERTIFICATIONS ──
    p = doc.add_paragraph()
    run = p.add_run("{%p if has_certs %}")
    run.font.size = Pt(1)
    run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

    _add_section_heading(doc, "Certifications")

    p = doc.add_paragraph()
    run = p.add_run("{%p for cert in certifications %}")
    run.font.size = Pt(1)
    run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run("{{cert.name}}")
    run.bold = True
    run.font.name = "Calibri"
    run.font.size = Pt(11)
    run2 = p.add_run("{% if cert.issuer %} \u2014 {{cert.issuer}}{% endif %}"
                      "{% if cert.date %}, {{cert.date}}{% endif %}")
    run2.font.name = "Calibri"
    run2.font.size = Pt(10)
    run2.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

    p = doc.add_paragraph()
    run = p.add_run("{%p endfor %}")
    run.font.size = Pt(1)
    run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

    p = doc.add_paragraph()
    run = p.add_run("{%p endif %}")
    run.font.size = Pt(1)
    run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

    # Save
    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(output_path))
    print(f"Template created: {output_path}")


if __name__ == "__main__":
    create_template(Path("templates/default.docx"))
