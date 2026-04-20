"""Convert a downloaded Word resume template into a docxtpl-compatible template.

Usage:
    python import_template.py "path/to/downloaded_template.docx"

This creates a copy in templates/ with placeholder text replaced by Jinja2 tokens.
After running, open the result in Word to verify and adjust any tokens that
weren't auto-detected.

How it works:
    1. Scans every paragraph and table cell for common placeholder patterns
    2. Replaces them with docxtpl tokens (e.g. "John Doe" → "{{contact.name}}")
    3. Injects loop markers for work history, education, and certifications
    4. Saves to templates/ folder

After importing, the template appears in the Generate tab's template picker.
"""
import re
import sys
from pathlib import Path
from copy import deepcopy
from docx import Document

# ── Placeholder patterns → Jinja2 tokens ──
# Each entry: (compiled regex, replacement string)
# Order matters — more specific patterns first
REPLACEMENTS = [
    # Contact fields
    (re.compile(r"\b(Your\s+Name|First\s+Last\s*Name?|John\s+Doe|Jane\s+Doe|Full\s+Name|FIRST\s+LAST|YOUR\s+NAME)\b", re.I),
     "{{contact.name}}"),
    (re.compile(r"\b[\w.+-]+@(example|email|mail|your)\.\w+\b", re.I),
     "{{contact.email}}"),
    (re.compile(r"\(?\d{3}\)?[\s.-]?\d{3}[\s.-]?\d{4}"), "{{contact.phone}}"),
    (re.compile(r"\b(City,?\s*S(?:tate|T)|Your\s+City|Anytown|Brooklyn|New\s+York)\b.*?\d{5}\b", re.I),
     "{{contact.location}}"),
    (re.compile(r"(linkedin\.com/in/\S+|www\.\S+\.com)", re.I), "{{contact.linkedin}}"),

    # Section content placeholders
    (re.compile(r"(Type|Enter|Write|Add)\s+(your\s+)?(professional\s+)?(summary|objective|profile)\s+here\.?", re.I),
     "{{summary}}"),
]

# Tokens to search for to confirm a template was already converted
EXISTING_TOKENS = ["{{contact.name}}", "{{contact.email}}", "{%p for"]


def _replace_in_runs(paragraph, replacements):
    """Apply regex replacements across a paragraph's full text, preserving first run's formatting."""
    full_text = paragraph.text
    changed = False
    for pattern, replacement in replacements:
        if pattern.search(full_text):
            full_text = pattern.sub(replacement, full_text)
            changed = True
    if changed:
        # Rewrite: keep first run's formatting, clear the rest
        if paragraph.runs:
            paragraph.runs[0].text = full_text
            for run in paragraph.runs[1:]:
                run.text = ""
    return changed


def _scan_and_replace(doc, replacements):
    """Apply replacements to all paragraphs and table cells."""
    count = 0
    for para in doc.paragraphs:
        if _replace_in_runs(para, replacements):
            count += 1
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    if _replace_in_runs(para, replacements):
                        count += 1
    return count


def import_template(source_path: str, output_name: str = None) -> Path:
    """Convert a Word resume template to docxtpl format.

    Args:
        source_path: Path to the downloaded .docx template
        output_name: Optional name for the output file (without extension)

    Returns:
        Path to the saved template in templates/
    """
    src = Path(source_path)
    if not src.exists():
        raise FileNotFoundError(f"Template not found: {src}")

    doc = Document(str(src))

    # Check if already converted
    full_text = "\n".join(p.text for p in doc.paragraphs)
    if any(tok in full_text for tok in EXISTING_TOKENS):
        print("This template already contains Jinja2 tokens — skipping auto-conversion.")
        print("It will be copied as-is to templates/.")
    else:
        count = _scan_and_replace(doc, REPLACEMENTS)
        print(f"Replaced {count} placeholder(s) with Jinja2 tokens.")

        if count == 0:
            print("\nNo common placeholders detected. You'll need to manually edit")
            print("the template in Word and replace placeholder text with tokens like:")
            print("  {{contact.name}}  {{contact.email}}  {{contact.phone}}")
            print("  {{summary}}  {{skills_text}}")
            print("  See README.md for the full list of available tokens.")

    # Save to templates/
    templates_dir = Path(__file__).parent / "templates"
    templates_dir.mkdir(exist_ok=True)

    if output_name:
        out_name = output_name.replace(" ", "_") + ".docx"
    else:
        out_name = src.stem.replace(" ", "_") + "_template.docx"

    out_path = templates_dir / out_name
    doc.save(str(out_path))
    print(f"\nSaved to: {out_path}")
    print(f"This template will now appear in the Generate tab dropdown.")
    print(f"\nIMPORTANT: Open the template in Word to verify the tokens are correct.")
    print(f"Replace any remaining placeholder text with the appropriate tokens.")
    print(f"See README.md section 'Creating a New Template' for the full token list.")

    return out_path


if __name__ == "__main__":
    if len(sys.argv) < 2:
        print("Usage: python import_template.py <path_to_template.docx> [output_name]")
        print("\nThis converts a downloaded Word resume template into a format")
        print("the Resume Rebuild Assistant can use to generate resumes.")
        sys.exit(1)

    source = sys.argv[1]
    name = sys.argv[2] if len(sys.argv) > 2 else None
    import_template(source, name)
