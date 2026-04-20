"""Load resume files into a normalized format for extraction.

Returns {"type": "text"|"images", "content": str|list[PIL.Image]}
"""
from pathlib import Path

from PIL import Image


def load_resume(filepath: str | Path) -> dict:
    """Dispatch by extension, return content ready for Claude extraction."""
    p = Path(filepath)
    ext = p.suffix.lower()

    if ext == ".docx":
        return _load_docx(p)
    elif ext == ".pdf":
        return _load_pdf(p)
    elif ext in (".png", ".jpg", ".jpeg"):
        return {"type": "images", "content": [Image.open(p).convert("RGB")]}
    else:
        raise ValueError(f"Unsupported file type: {ext}")


def _load_docx(path: Path) -> dict:
    from docx import Document

    doc = Document(str(path))
    paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
    # Also grab table cell text (some resumes use tables for layout)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                txt = cell.text.strip()
                if txt:
                    paragraphs.append(txt)
    return {"type": "text", "content": "\n".join(paragraphs)}


def _load_pdf(path: Path) -> dict:
    from pypdf import PdfReader

    reader = PdfReader(str(path))
    text_parts = []
    for page in reader.pages:
        t = page.extract_text()
        if t and t.strip():
            text_parts.append(t.strip())

    full_text = "\n\n".join(text_parts)

    # If pypdf extracted meaningful text, use it (skip image conversion)
    if len(full_text) > 50:
        return {"type": "text", "content": full_text}

    # Scanned PDF — convert pages to images
    try:
        from pdf2image import convert_from_path
        from config import SCAN_DPI
        images = convert_from_path(str(path), dpi=SCAN_DPI)
        return {"type": "images", "content": images}
    except Exception as e:
        err = str(e).lower()
        if "poppler" in err or "pdftoppm" in err:
            raise RuntimeError(
                f"This PDF appears to be a scan (no extractable text layer). "
                f"Converting scans to images requires Poppler.\n\n"
                f"Install: download from https://github.com/osber/poppler-windows/releases , "
                f"extract, and add the bin\\ folder to your system PATH. Then restart the app.\n\n"
                f"Original error: {e}"
            ) from e
        raise
