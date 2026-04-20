"""Render structured resume data to Word and PDF output files."""
import copy
from pathlib import Path

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docxtpl import DocxTemplate

# ── Default format settings ──
DEFAULT_FORMAT = {
    "font_family": "Calibri",
    "name_size": 18,        # pt
    "contact_size": 10,
    "section_header_size": 12,
    "body_size": 10.5,
    "date_size": 10,
    "line_spacing": 1.0,
    "margin_top": 0.5,      # inches
    "margin_bottom": 0.5,
    "margin_left": 0.65,
    "margin_right": 0.65,
    "section_spacing_before": 10,  # pt before section headers
    "section_spacing_after": 3,
    "paragraph_spacing": 1,  # pt between body paragraphs
    "bullet_indent": 0.25,  # inches
    "show_section_lines": True,
    "name_color": "#1a1a1a",
    "header_color": "#1a1a2e",
    "body_color": "#1a1a1a",
    "accent_color": "#444444",
    "section_line_color": "#999999",
    "section_order": ["summary", "skills", "experience", "education", "certifications"],
    "contact_layout": "center",  # "center", "right", or "left"
    # Two-column sidebar layout settings
    "layout_style": "single",  # "single" or "two-column"
    "header_banner_color": "#d4c5a9",  # beige banner behind name
    "tagline": "",  # subtitle under name (e.g. "ATTORNEY", "SOFTWARE ENGINEER")
    "sidebar_sections": ["education", "skills", "certifications"],  # sections in left column
    "sidebar_width_pct": 35,  # sidebar width as percentage of page
}


def _hex_to_rgb(hex_color: str) -> RGBColor:
    """Convert '#RRGGBB' to a python-docx RGBColor."""
    h = hex_color.lstrip("#")
    return RGBColor(int(h[0:2], 16), int(h[2:4], 16), int(h[4:6], 16))


def _set_run_font(run, font_family: str, size_pt: float, color_hex: str,
                  bold: bool = False, italic: bool = False):
    """Apply font settings to a run explicitly (no style inheritance)."""
    run.font.name = font_family
    run.font.size = Pt(size_pt)
    run.font.color.rgb = _hex_to_rgb(color_hex)
    run.font.bold = bold
    run.font.italic = italic
    # Force font for East Asian / complex script fallback
    r_elem = run._element
    rPr = r_elem.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = r_elem.makeelement(qn("w:rFonts"), {})
        rPr.insert(0, rFonts)
    rFonts.set(qn("w:ascii"), font_family)
    rFonts.set(qn("w:hAnsi"), font_family)
    rFonts.set(qn("w:cs"), font_family)


def _set_paragraph_spacing(para, space_before_pt: float = 0, space_after_pt: float = 0,
                           line_spacing: float = 1.0):
    """Set paragraph spacing attributes."""
    pf = para.paragraph_format
    pf.space_before = Pt(space_before_pt)
    pf.space_after = Pt(space_after_pt)
    pf.line_spacing = line_spacing


def _add_section_border(para, color_hex: str):
    """Add a thin bottom border to a paragraph (for section header lines)."""
    pPr = para._element.get_or_add_pPr()
    pBdr = pPr.find(qn("w:pBdr"))
    if pBdr is None:
        pBdr = para._element.makeelement(qn("w:pBdr"), {})
        pPr.append(pBdr)
    bottom = para._element.makeelement(qn("w:bottom"), {
        qn("w:val"): "single",
        qn("w:sz"): "4",
        qn("w:space"): "1",
        qn("w:color"): color_hex.lstrip("#"),
    })
    pBdr.append(bottom)


# ══════════════════════════════════════════════
#  Template-based rendering (backward compat)
# ══════════════════════════════════════════════

def _prepare_context(data: dict) -> dict:
    """Pre-process resume data into template-friendly context."""
    ctx = {
        "contact": data.get("contact", {}),
        "summary": data.get("summary", ""),
        "work_history": data.get("work_history", []),
        "education": data.get("education", []),
        "certifications": data.get("certifications", []),
        "skills_text": " \u2022 ".join(data.get("skills", [])),
        "has_summary": bool(data.get("summary", "").strip()),
        "has_skills": bool(data.get("skills")),
        "has_work": bool(data.get("work_history")),
        "has_education": bool(data.get("education")),
        "has_certs": bool(data.get("certifications")),
    }
    parts = []
    c = ctx["contact"]
    if c.get("email"):
        parts.append(c["email"])
    if c.get("phone"):
        parts.append(c["phone"])
    if c.get("location"):
        parts.append(c["location"])
    ctx["contact_line"] = " | ".join(parts)
    return ctx


def render_docx(data: dict, template_path: str | Path, output_path: str | Path) -> Path:
    """Fill a docxtpl Word template with resume data and save."""
    tpl = DocxTemplate(str(template_path))
    ctx = _prepare_context(data)
    tpl.render(ctx)
    out = Path(output_path)
    out.parent.mkdir(parents=True, exist_ok=True)
    tpl.save(str(out))
    return out


def _convert_docx_to_pdf(docx_path: str, pdf_path: str):
    """Convert a docx to PDF. Uses Word COM on Windows, raises clear error on Linux/cloud."""
    import platform
    if platform.system() == "Windows":
        try:
            import pythoncom
            pythoncom.CoInitialize()
            try:
                from docx2pdf import convert
                convert(docx_path, pdf_path)
            finally:
                pythoncom.CoUninitialize()
        except ImportError:
            raise RuntimeError(
                "PDF conversion requires docx2pdf and pywin32. "
                "Install with: pip install docx2pdf pywin32"
            )
    else:
        # Linux/Mac/Cloud — try libreoffice as fallback
        import subprocess
        import shutil
        lo = shutil.which("libreoffice") or shutil.which("soffice")
        if lo:
            subprocess.run([lo, "--headless", "--convert-to", "pdf",
                           "--outdir", str(Path(pdf_path).parent), docx_path],
                          check=True, capture_output=True)
            # libreoffice outputs to same dir with .pdf extension
            expected = Path(docx_path).with_suffix(".pdf")
            if expected != Path(pdf_path) and expected.exists():
                expected.rename(pdf_path)
        else:
            raise RuntimeError(
                "PDF export is not available in the cloud. "
                "Download the Word document instead and convert to PDF locally."
            )


def render_pdf(data: dict, template_path: str | Path, output_path: str | Path) -> Path:
    """Render resume to PDF by first creating a docx then converting via Word."""
    out_pdf = Path(output_path)
    tmp_docx = out_pdf.with_suffix(".tmp.docx")
    render_docx(data, template_path, tmp_docx)
    _convert_docx_to_pdf(str(tmp_docx), str(out_pdf))
    tmp_docx.unlink(missing_ok=True)
    return out_pdf


# ══════════════════════════════════════════════
#  Styled rendering (programmatic, full control)
# ══════════════════════════════════════════════

def render_docx_styled(data: dict, fmt: dict, output_path: str | Path) -> Path:
    """Build a polished, ATS-friendly resume from scratch with python-docx."""
    f = {**DEFAULT_FORMAT, **(fmt or {})}

    if f.get("layout_style") == "two-column":
        return _build_two_column(data, f, output_path)

    doc = Document()

    # ── Page margins ──
    for section in doc.sections:
        section.top_margin = Inches(f["margin_top"])
        section.bottom_margin = Inches(f["margin_bottom"])
        section.left_margin = Inches(f["margin_left"])
        section.right_margin = Inches(f["margin_right"])

    contact = data.get("contact", {})
    font = f["font_family"]
    ls = f["line_spacing"]

    # ── Helper: add a section header ──
    def add_section_header(title: str):
        para = doc.add_paragraph()
        para.alignment = WD_ALIGN_PARAGRAPH.LEFT
        run = para.add_run(title.upper())
        _set_run_font(run, font, f["section_header_size"], f["header_color"], bold=True)
        _set_paragraph_spacing(para, f["section_spacing_before"], f["section_spacing_after"], ls)
        para.paragraph_format.keep_with_next = True
        if f["show_section_lines"]:
            _add_section_border(para, f["section_line_color"])
        return para

    # ── Helper: add body paragraph ──
    def add_body(text: str, bold: bool = False, italic: bool = False,
                 color: str | None = None, size: float | None = None,
                 alignment=WD_ALIGN_PARAGRAPH.LEFT,
                 space_before: float = 0, space_after: float = 1):
        para = doc.add_paragraph()
        para.alignment = alignment
        run = para.add_run(text)
        _set_run_font(run, font, size or f["body_size"],
                      color or f["body_color"], bold=bold, italic=italic)
        _set_paragraph_spacing(para, space_before, space_after, ls)
        return para

    # ── Helper: add a mixed-format line (bold part + normal part) ──
    def add_mixed_line(bold_text: str, normal_text: str,
                       size: float | None = None, color: str | None = None,
                       space_before: float = 0, space_after: float = 1):
        para = doc.add_paragraph()
        para.alignment = WD_ALIGN_PARAGRAPH.LEFT
        sz = size or f["body_size"]
        clr = color or f["body_color"]
        if bold_text:
            run_b = para.add_run(bold_text)
            _set_run_font(run_b, font, sz, clr, bold=True)
        if normal_text:
            run_n = para.add_run(normal_text)
            _set_run_font(run_n, font, sz, clr)
        _set_paragraph_spacing(para, space_before, space_after, ls)
        return para

    # ═══════════════════════════════
    #  NAME BLOCK
    # ═══════════════════════════════
    layout = f.get("contact_layout", "center")
    name_align = {
        "center": WD_ALIGN_PARAGRAPH.CENTER,
        "left": WD_ALIGN_PARAGRAPH.LEFT,
        "right": WD_ALIGN_PARAGRAPH.RIGHT,
    }.get(layout, WD_ALIGN_PARAGRAPH.CENTER)

    name_text = contact.get("name", "").strip()

    if layout == "right":
        # Right-aligned contact: Name bold on the left, contact info on the right
        # Use a two-column table with invisible borders
        table = doc.add_table(rows=1, cols=2)
        table.autofit = True
        # Remove table borders
        tbl = table._element
        tblPr = tbl.find(qn("w:tblPr"))
        if tblPr is None:
            tblPr = tbl.makeelement(qn("w:tblPr"), {})
            tbl.insert(0, tblPr)
        tblBorders = tblPr.makeelement(qn("w:tblBorders"), {})
        for edge in ["top", "left", "bottom", "right", "insideH", "insideV"]:
            border_el = tblPr.makeelement(qn(f"w:{edge}"), {
                qn("w:val"): "none", qn("w:sz"): "0", qn("w:space"): "0", qn("w:color"): "auto",
            })
            tblBorders.append(border_el)
        tblPr.append(tblBorders)

        # Left cell: name
        left_cell = table.cell(0, 0)
        left_cell.text = ""
        p = left_cell.paragraphs[0]
        run = p.add_run(name_text)
        _set_run_font(run, font, f["name_size"], f["name_color"], bold=True)
        _set_paragraph_spacing(p, 0, 0, ls)

        # Right cell: contact info stacked
        right_cell = table.cell(0, 1)
        right_cell.text = ""
        contact_parts = []
        if contact.get("email"):
            contact_parts.append(contact["email"])
        if contact.get("phone"):
            contact_parts.append(contact["phone"])
        if contact.get("location"):
            contact_parts.append(contact["location"])
        linkedin = contact.get("linkedin", "").strip()
        if linkedin:
            contact_parts.append(linkedin)
        p = right_cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        run = p.add_run("\n".join(contact_parts))
        _set_run_font(run, font, f["contact_size"], f["accent_color"])
        _set_paragraph_spacing(p, 0, 4, ls)
    else:
        # Center or left layout
        if name_text:
            para = doc.add_paragraph()
            para.alignment = name_align
            run = para.add_run(name_text)
            _set_run_font(run, font, f["name_size"], f["name_color"], bold=True)
            _set_paragraph_spacing(para, 0, 0, ls)

        contact_parts = []
        if contact.get("email"):
            contact_parts.append(contact["email"])
        if contact.get("phone"):
            contact_parts.append(contact["phone"])
        if contact.get("location"):
            contact_parts.append(contact["location"])
        if contact_parts:
            para = doc.add_paragraph()
            para.alignment = name_align
            run = para.add_run(" | ".join(contact_parts))
            _set_run_font(run, font, f["contact_size"], f["accent_color"])
            _set_paragraph_spacing(para, 0, 2, ls)

        linkedin = contact.get("linkedin", "").strip()
        if linkedin:
            para = doc.add_paragraph()
            para.alignment = name_align
            run = para.add_run(linkedin)
            _set_run_font(run, font, f["contact_size"], f["accent_color"])
            _set_paragraph_spacing(para, 0, 4, ls)

    # ═══════════════════════════════
    #  SECTION BUILDERS (called by order)
    # ═══════════════════════════════
    def _build_summary():
        summary = data.get("summary", "").strip()
        if not summary:
            return
        add_section_header("Summary")
        add_body(summary, space_after=f["paragraph_spacing"])

    def _build_skills():
        skills = data.get("skills", [])
        if not skills:
            return
        add_section_header("Skills")
        add_body(" \u2022 ".join(skills), space_after=f["paragraph_spacing"])

    def _build_experience():
        work = data.get("work_history", [])
        if not work:
            return
        add_section_header("Experience")
        for job in work:
            role = job.get("role", "").strip()
            employer = job.get("employer", "").strip()

            if role and role != "Role" and employer:
                add_mixed_line(role, f" \u2014 {employer}", space_before=6, space_after=0)
            elif role and role != "Role":
                add_body(role, bold=True, space_before=6, space_after=0)
            elif employer:
                add_body(employer, bold=True, space_before=6, space_after=0)

            date_parts = []
            start = job.get("start_date", "").strip()
            end = job.get("end_date", "").strip()
            if start or end:
                date_parts.append(f"{start} \u2013 {end}" if start and end else (start or end))
            loc = job.get("location", "").strip()
            if loc:
                date_parts.append(loc)
            if date_parts:
                add_body(" | ".join(date_parts), italic=True,
                         size=f["date_size"], color=f["accent_color"],
                         space_before=0, space_after=1)

            for b in job.get("bullets", []):
                b_text = b.strip()
                if not b_text:
                    continue
                para = doc.add_paragraph(style="List Bullet")
                para.paragraph_format.left_indent = Inches(f["bullet_indent"] + 0.2)
                para.paragraph_format.first_line_indent = Inches(-0.2)
                _set_paragraph_spacing(para, 0, f["paragraph_spacing"], ls)
                run = para.add_run(b_text)
                _set_run_font(run, font, f["body_size"], f["body_color"])

    def _build_education():
        edu = data.get("education", [])
        if not edu:
            return
        add_section_header("Education")
        for entry in edu:
            institution = entry.get("institution", "").strip()
            degree = entry.get("degree", "").strip()
            field = entry.get("field", "").strip()
            year = entry.get("graduation_year", "").strip()
            honors = entry.get("honors", "").strip()

            normal_parts = []
            if degree and field:
                normal_parts.append(f"{degree} in {field}")
            elif degree:
                normal_parts.append(degree)
            elif field:
                normal_parts.append(field)
            if year:
                normal_parts.append(year)

            normal_text = ""
            if normal_parts:
                normal_text = " \u2014 " + ", ".join(normal_parts)

            if institution or normal_text:
                add_mixed_line(institution, normal_text, space_before=4, space_after=1)

            if honors:
                add_body(honors, italic=True, size=f["date_size"],
                         color=f["accent_color"], space_before=0, space_after=2)

    def _build_certifications():
        certs = data.get("certifications", [])
        if not certs:
            return
        add_section_header("Certifications")
        for cert in certs:
            cert_name = cert.get("name", "").strip()
            issuer = cert.get("issuer", "").strip()
            date = cert.get("date", "").strip()

            normal_parts = []
            if issuer:
                normal_parts.append(issuer)
            if date:
                normal_parts.append(date)

            normal_text = ""
            if normal_parts:
                normal_text = " \u2014 " + ", ".join(normal_parts)

            if cert_name or normal_text:
                add_mixed_line(cert_name, normal_text,
                               size=f["date_size"], space_before=2, space_after=1)

    # ── Render sections in configured order ──
    section_builders = {
        "summary": _build_summary,
        "skills": _build_skills,
        "experience": _build_experience,
        "education": _build_education,
        "certifications": _build_certifications,
    }
    for section_key in f.get("section_order", DEFAULT_FORMAT["section_order"]):
        builder = section_builders.get(section_key)
        if builder:
            builder()

    # ── Save ──
    out = Path(output_path)
    out.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(out))
    return out


def _set_cell_shading(cell, hex_color: str):
    """Set background shading on a table cell."""
    tcPr = cell._element.get_or_add_tcPr()
    shd = tcPr.find(qn("w:shd"))
    if shd is not None:
        tcPr.remove(shd)
    shd = cell._element.makeelement(qn("w:shd"), {
        qn("w:fill"): hex_color.lstrip("#"),
        qn("w:val"): "clear",
        qn("w:color"): "auto",
    })
    tcPr.append(shd)


def _remove_table_borders(table):
    """Remove all borders from a Word table."""
    tbl = table._element
    tblPr = tbl.find(qn("w:tblPr"))
    if tblPr is None:
        tblPr = tbl.makeelement(qn("w:tblPr"), {})
        tbl.insert(0, tblPr)
    # Remove existing borders
    existing = tblPr.find(qn("w:tblBorders"))
    if existing is not None:
        tblPr.remove(existing)
    tblBorders = tblPr.makeelement(qn("w:tblBorders"), {})
    for edge in ["top", "left", "bottom", "right", "insideH", "insideV"]:
        border_el = tblPr.makeelement(qn(f"w:{edge}"), {
            qn("w:val"): "none", qn("w:sz"): "0",
            qn("w:space"): "0", qn("w:color"): "auto",
        })
        tblBorders.append(border_el)
    tblPr.append(tblBorders)


def _set_cell_width(cell, inches_val):
    """Set preferred width on a cell."""
    tcPr = cell._element.get_or_add_tcPr()
    tcW = tcPr.find(qn("w:tcW"))
    if tcW is None:
        tcW = cell._element.makeelement(qn("w:tcW"), {})
        tcPr.append(tcW)
    tcW.set(qn("w:w"), str(int(inches_val * 1440)))  # 1440 twips per inch
    tcW.set(qn("w:type"), "dxa")


def _set_cell_margins(cell, top=0, left=100, bottom=0, right=100):
    """Set cell margins in twips."""
    tcPr = cell._element.get_or_add_tcPr()
    tcMar = tcPr.find(qn("w:tcMar"))
    if tcMar is not None:
        tcPr.remove(tcMar)
    tcMar = cell._element.makeelement(qn("w:tcMar"), {})
    for edge, val in [("top", top), ("left", left), ("bottom", bottom), ("right", right)]:
        el = cell._element.makeelement(qn(f"w:{edge}"), {
            qn("w:w"): str(val), qn("w:type"): "dxa",
        })
        tcMar.append(el)
    tcPr.append(tcMar)


def _add_cell_paragraph(cell, text, font_family, size, color_hex,
                        bold=False, italic=False, space_before=0,
                        space_after=2, line_spacing=1.0, alignment=None):
    """Add a formatted paragraph to a table cell."""
    # Use existing first paragraph if cell is empty, else add new
    if cell.paragraphs and cell.paragraphs[-1].text == "" and len(cell.paragraphs) <= 1:
        para = cell.paragraphs[0]
    else:
        para = cell.add_paragraph()
    if alignment is not None:
        para.alignment = alignment
    run = para.add_run(text)
    _set_run_font(run, font_family, size, color_hex, bold=bold, italic=italic)
    _set_paragraph_spacing(para, space_before, space_after, line_spacing)
    return para


def _add_cell_section_header(cell, title, f, show_line=True):
    """Add a section header inside a table cell."""
    para = cell.add_paragraph()
    run = para.add_run(title.upper())
    _set_run_font(run, f["font_family"], f["section_header_size"],
                  f["header_color"], bold=True)
    _set_paragraph_spacing(para, f["section_spacing_before"],
                           f["section_spacing_after"], f["line_spacing"])
    if show_line and f.get("show_section_lines", True):
        _add_section_border(para, f["section_line_color"])
    return para


def _build_two_column(data: dict, f: dict, output_path) -> Path:
    """Build a two-column sidebar resume layout (like the Alta Parks template)."""
    doc = Document()
    font = f["font_family"]
    ls = f["line_spacing"]
    contact = data.get("contact", {})

    # Page margins — tighter for two-column
    for section in doc.sections:
        section.top_margin = Inches(0)  # banner goes to edge
        section.bottom_margin = Inches(f["margin_bottom"])
        section.left_margin = Inches(f["margin_left"])
        section.right_margin = Inches(f["margin_right"])

    page_width = 8.5 - f["margin_left"] - f["margin_right"]
    sidebar_w = page_width * f.get("sidebar_width_pct", 35) / 100
    main_w = page_width - sidebar_w

    # ═══════════════════════════════
    #  HEADER BANNER
    # ═══════════════════════════════
    banner_color = f.get("header_banner_color", "#d4c5a9")
    header_table = doc.add_table(rows=1, cols=1)
    _remove_table_borders(header_table)
    header_table.autofit = False
    hcell = header_table.cell(0, 0)
    _set_cell_shading(hcell, banner_color)
    _set_cell_margins(hcell, top=300, left=200, bottom=300, right=200)
    _set_cell_width(hcell, page_width)

    # Name
    name_text = contact.get("name", "").strip()
    if name_text:
        p = hcell.paragraphs[0]
        run = p.add_run(name_text.upper())
        # Use a larger, more dramatic name size for two-column
        name_sz = f["name_size"] + 8
        _set_run_font(run, font, name_sz, f["name_color"], bold=True)
        _set_paragraph_spacing(p, 0, 0, ls)
        # Add letter spacing for the dramatic header look
        rPr = run._element.get_or_add_rPr()
        spacing_el = rPr.find(qn("w:spacing"))
        if spacing_el is None:
            spacing_el = run._element.makeelement(qn("w:spacing"), {})
            rPr.append(spacing_el)
        spacing_el.set(qn("w:val"), "60")  # 3pt letter spacing

    # Tagline (e.g. "ATTORNEY")
    tagline = f.get("tagline", "").strip()
    if tagline:
        p = hcell.add_paragraph()
        run = p.add_run(tagline.upper())
        _set_run_font(run, font, f["section_header_size"], f["body_color"], bold=False)
        _set_paragraph_spacing(p, 4, 0, ls)
        # Letter spacing on tagline too
        rPr = run._element.get_or_add_rPr()
        spacing_el = run._element.makeelement(qn("w:spacing"), {})
        spacing_el.set(qn("w:val"), "40")
        rPr.append(spacing_el)

    # ═══════════════════════════════
    #  TWO-COLUMN BODY
    # ═══════════════════════════════
    body_table = doc.add_table(rows=1, cols=2)
    _remove_table_borders(body_table)
    body_table.autofit = False

    sidebar = body_table.cell(0, 0)
    main = body_table.cell(0, 1)
    _set_cell_width(sidebar, sidebar_w)
    _set_cell_width(main, main_w)
    _set_cell_margins(sidebar, top=150, left=50, bottom=100, right=150)
    _set_cell_margins(main, top=150, left=150, bottom=100, right=50)

    # Clear default paragraphs
    sidebar.paragraphs[0].text = ""
    main.paragraphs[0].text = ""

    # ── SIDEBAR: Contact info (always first) ──
    _add_cell_section_header(sidebar, "Contact", f)

    contact_fields = []
    if contact.get("location"):
        contact_fields.append(contact["location"])
    if contact.get("phone"):
        contact_fields.append(contact["phone"])
    if contact.get("email"):
        contact_fields.append(contact["email"])
    if contact.get("linkedin"):
        contact_fields.append(contact["linkedin"])

    for cf in contact_fields:
        _add_cell_paragraph(sidebar, cf, font, f["body_size"],
                            f["body_color"], space_after=1, line_spacing=ls)

    # ── SIDEBAR: Configured sections ──
    sidebar_keys = f.get("sidebar_sections", ["education", "skills", "certifications"])

    def _sidebar_education():
        edu = data.get("education", [])
        if not edu:
            return
        _add_cell_section_header(sidebar, "Education", f)
        for entry in edu:
            degree = entry.get("degree", "").strip()
            field = entry.get("field", "").strip()
            year = entry.get("graduation_year", "").strip()
            institution = entry.get("institution", "").strip()
            honors = entry.get("honors", "").strip()

            # Degree line
            degree_text = ""
            if degree and field:
                degree_text = f"{degree.upper()} IN {field.upper()}"
            elif degree:
                degree_text = degree.upper()
            if year:
                degree_text += f" \u2022 {year}" if degree_text else year

            if degree_text:
                _add_cell_paragraph(sidebar, degree_text, font, f["date_size"],
                                    f["header_color"], bold=True, space_before=6, space_after=1, line_spacing=ls)
            if institution:
                _add_cell_paragraph(sidebar, institution, font, f["body_size"],
                                    f["body_color"], space_after=1, line_spacing=ls)
            if honors:
                _add_cell_paragraph(sidebar, honors, font, f["date_size"],
                                    f["accent_color"], italic=True, space_after=3, line_spacing=ls)

    def _sidebar_skills():
        skills = data.get("skills", [])
        if not skills:
            return
        _add_cell_section_header(sidebar, "Skills", f)
        for skill in skills:
            _add_cell_paragraph(sidebar, skill, font, f["body_size"],
                                f["body_color"], space_after=1, line_spacing=ls)

    def _sidebar_certifications():
        certs = data.get("certifications", [])
        if not certs:
            return
        _add_cell_section_header(sidebar, "Certifications", f)
        for cert in certs:
            cert_name = cert.get("name", "").strip()
            detail_parts = []
            if cert.get("issuer", "").strip():
                detail_parts.append(cert["issuer"].strip())
            if cert.get("date", "").strip():
                detail_parts.append(cert["date"].strip())
            if cert_name:
                _add_cell_paragraph(sidebar, cert_name, font, f["body_size"],
                                    f["body_color"], bold=True, space_before=3, space_after=0, line_spacing=ls)
            if detail_parts:
                _add_cell_paragraph(sidebar, ", ".join(detail_parts), font, f["date_size"],
                                    f["accent_color"], space_after=2, line_spacing=ls)

    def _sidebar_summary():
        summary = data.get("summary", "").strip()
        if not summary:
            return
        _add_cell_section_header(sidebar, "Profile", f)
        _add_cell_paragraph(sidebar, summary, font, f["body_size"],
                            f["body_color"], space_after=3, line_spacing=ls)

    sidebar_builders = {
        "education": _sidebar_education,
        "skills": _sidebar_skills,
        "certifications": _sidebar_certifications,
        "summary": _sidebar_summary,
    }

    for sk in sidebar_keys:
        builder = sidebar_builders.get(sk)
        if builder:
            builder()

    # ── MAIN COLUMN: Sections not in sidebar ──
    main_sections = [s for s in f.get("section_order", DEFAULT_FORMAT["section_order"])
                     if s not in sidebar_keys]

    def _main_summary():
        summary = data.get("summary", "").strip()
        if not summary:
            return
        _add_cell_section_header(main, "Profile", f)
        _add_cell_paragraph(main, summary, font, f["body_size"],
                            f["body_color"], space_after=4, line_spacing=ls)

    def _main_experience():
        work = data.get("work_history", [])
        if not work:
            return
        _add_cell_section_header(main, "Experience", f)
        for job in work:
            role = job.get("role", "").strip()
            employer = job.get("employer", "").strip()
            start = job.get("start_date", "").strip()
            end = job.get("end_date", "").strip()
            loc = job.get("location", "").strip()

            # Title line: ROLE • START—END
            title_parts = []
            if role and role != "Role":
                title_parts.append(role.upper())
            if start or end:
                date_str = f"{start}\u2014{end}" if start and end else (start or end)
                title_parts.append(date_str)
            if title_parts:
                _add_cell_paragraph(main, " \u2022 ".join(title_parts), font,
                                    f["date_size"], f["header_color"], bold=True,
                                    space_before=8, space_after=1, line_spacing=ls)

            # Employer line
            emp_parts = []
            if employer:
                emp_parts.append(employer)
            if loc:
                emp_parts.append(loc)
            if emp_parts:
                _add_cell_paragraph(main, " \u2022 ".join(emp_parts), font,
                                    f["body_size"], f["body_color"],
                                    space_after=3, line_spacing=ls)

            # Bullets — rendered as proper bullet list inside the cell
            bullets = [b.strip() for b in job.get("bullets", []) if b.strip()]
            for b_text in bullets:
                bp = main.add_paragraph(style="List Bullet")
                bp.paragraph_format.left_indent = Inches(f["bullet_indent"] + 0.2)
                bp.paragraph_format.first_line_indent = Inches(-0.2)
                _set_paragraph_spacing(bp, 0, f["paragraph_spacing"], ls)
                run = bp.add_run(b_text)
                _set_run_font(run, font, f["body_size"], f["body_color"])

    def _main_education():
        edu = data.get("education", [])
        if not edu:
            return
        _add_cell_section_header(main, "Education", f)
        for entry in edu:
            institution = entry.get("institution", "").strip()
            degree = entry.get("degree", "").strip()
            field = entry.get("field", "").strip()
            year = entry.get("graduation_year", "").strip()
            honors = entry.get("honors", "").strip()

            parts = []
            if degree:
                parts.append(degree)
            if field:
                parts.append(f"in {field}")
            if year:
                parts.append(f"({year})")
            line1 = " ".join(parts)

            if institution:
                p = main.add_paragraph()
                run_b = p.add_run(institution)
                _set_run_font(run_b, font, f["body_size"], f["body_color"], bold=True)
                if line1:
                    run_n = p.add_run(f" \u2014 {line1}")
                    _set_run_font(run_n, font, f["body_size"], f["body_color"])
                _set_paragraph_spacing(p, 4, 1, ls)

            if honors:
                _add_cell_paragraph(main, honors, font, f["date_size"],
                                    f["accent_color"], italic=True, space_after=2, line_spacing=ls)

    def _main_skills():
        skills = data.get("skills", [])
        if not skills:
            return
        _add_cell_section_header(main, "Skills", f)
        _add_cell_paragraph(main, " \u2022 ".join(skills), font, f["body_size"],
                            f["body_color"], space_after=3, line_spacing=ls)

    def _main_certifications():
        certs = data.get("certifications", [])
        if not certs:
            return
        _add_cell_section_header(main, "Certifications", f)
        for cert in certs:
            cert_name = cert.get("name", "").strip()
            parts = []
            if cert.get("issuer", "").strip():
                parts.append(cert["issuer"].strip())
            if cert.get("date", "").strip():
                parts.append(cert["date"].strip())
            p = main.add_paragraph()
            if cert_name:
                run_b = p.add_run(cert_name)
                _set_run_font(run_b, font, f["body_size"], f["body_color"], bold=True)
            if parts:
                run_n = p.add_run(f" \u2014 {', '.join(parts)}")
                _set_run_font(run_n, font, f["date_size"], f["accent_color"])
            _set_paragraph_spacing(p, 2, 1, ls)

    main_builders = {
        "summary": _main_summary,
        "experience": _main_experience,
        "education": _main_education,
        "skills": _main_skills,
        "certifications": _main_certifications,
    }

    for sk in main_sections:
        builder = main_builders.get(sk)
        if builder:
            builder()

    # ── Save ──
    out = Path(output_path)
    out.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(out))
    return out


def render_pdf_styled(data: dict, fmt: dict, output_path: str | Path) -> Path:
    """Render styled resume to PDF via temp docx + docx2pdf conversion."""
    out_pdf = Path(output_path)
    tmp_docx = out_pdf.with_suffix(".tmp.docx")
    render_docx_styled(data, fmt, tmp_docx)
    _convert_docx_to_pdf(str(tmp_docx), str(out_pdf))
    tmp_docx.unlink(missing_ok=True)
    return out_pdf
